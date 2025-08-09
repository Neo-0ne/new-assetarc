
from docx import Document

def fill_docx_template(template_path: str, values: dict, output_path: str):
    doc=Document(template_path)
    # paragraphs
    for p in doc.paragraphs:
        for k,v in values.items():
            if f'{{{{{k}}}}}' in p.text:
                for r in p.runs:
                    r.text=r.text.replace(f'{{{{{k}}}}}', str(v))
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k,v in values.items():
                    if f'{{{{{k}}}}}' in cell.text:
                        cell.text=cell.text.replace(f'{{{{{k}}}}}', str(v))
    doc.save(output_path)
