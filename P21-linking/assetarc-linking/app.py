
import os, tempfile, datetime, json
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from pydantic import BaseModel, ValidationError
from jinja2 import Template
from weasyprint import HTML
from docx import Document

load_dotenv()
app=Flask(__name__)
CORS(app, resources={r"/*": {"origins":[o.strip() for o in os.getenv('CORS_ALLOWED_ORIGINS','').split(',') if o.strip()]}},
     supports_credentials=True)

@app.get('/healthz')
def h(): return jsonify({'ok':True})

class GenBody(BaseModel):
    template_id: str
    values: dict
    output_name: str

@app.post('/generate/docx')
def gen_docx():
    try: b=GenBody(**request.get_json(force=True))
    except ValidationError as e: return jsonify({'ok':False,'error':e.errors()}),400
    tpl=os.path.join('templates', b.template_id)
    if not tpl.endswith('.docx') or not os.path.exists(tpl): return jsonify({'ok':False,'error':'docx template not found'}),404
    out=os.path.join(tempfile.gettempdir(), b.output_name if b.output_name.endswith('.docx') else b.output_name+'.docx')
    doc=Document(tpl)
    for p in doc.paragraphs:
        for k,v in b.values.items():
            if f'{{{k}}}' in p.text:
                for r in p.runs: r.text=r.text.replace(f'{{{k}}}', str(v))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k,v in b.values.items():
                    if f'{{{k}}}' in cell.text:
                        cell.text=cell.text.replace(f'{{{k}}}', str(v))
    doc.save(out); return send_file(out, as_attachment=True, download_name=os.path.basename(out))

@app.post('/generate/pdf')
def gen_pdf():
    try: b=GenBody(**request.get_json(force=True))
    except ValidationError as e: return jsonify({'ok':False,'error':e.errors()}),400
    tpl=os.path.join('templates', b.template_id)
    if not tpl.endswith('.html') or not os.path.exists(tpl): return jsonify({'ok':False,'error':'html template not found'}),404
    out=os.path.join(tempfile.gettempdir(), b.output_name if b.output_name.endswith('.pdf') else b.output_name+'.pdf')
    with open(tpl,'r',encoding='utf-8') as f: html=Template(f.read()).render(**b.values)
    HTML(string=html).write_pdf(out); return send_file(out, as_attachment=True, download_name=os.path.basename(out))

if __name__=='__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT','0') or '0'), debug=os.getenv('FLASK_ENV')=='development')
