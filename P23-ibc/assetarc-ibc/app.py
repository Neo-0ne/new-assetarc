
import os, tempfile, json
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from jinja2 import Template
from docx import Document
from weasyprint import HTML

load_dotenv()
app=Flask(__name__)
CORS(app, resources={r"/*": {"origins":[o.strip() for o in os.getenv('CORS_ALLOWED_ORIGINS','').split(',') if o.strip()]}}, supports_credentials=True)

@app.get('/healthz')
def h(): return jsonify({'ok':True})

@app.get('/templates')
def list_templates():
    files=[]
    for n in os.listdir('templates'):
        if n.endswith('.docx') or n.endswith('.html'):
            files.append(n)
    return jsonify({'ok':True,'templates':files})

@app.post('/generate/docx')
def gen_docx():
    b=request.get_json(force=True) or {}
    tid=b.get('template_id'); values=b.get('values',{}); outn=b.get('output_name','out.docx')
    tpl=os.path.join('templates', tid or '')
    if not tpl.endswith('.docx') or not os.path.exists(tpl): return jsonify({'ok':False,'error':'docx template not found'}),404
    out=os.path.join(tempfile.gettempdir(), outn if outn.endswith('.docx') else outn+'.docx')
    doc=Document(tpl)
    for p in doc.paragraphs:
        for k,v in values.items():
            if f'{{{k}}}' in p.text:
                for r in p.runs: r.text=r.text.replace(f'{{{k}}}', str(v))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for k,v in values.items():
                    if f'{{{k}}}' in cell.text:
                        cell.text=cell.text.replace(f'{{{k}}}', str(v))
    doc.save(out); return send_file(out, as_attachment=True, download_name=os.path.basename(out))

@app.post('/generate/pdf')
def gen_pdf():
    b=request.get_json(force=True) or {}
    tid=b.get('template_id'); values=b.get('values',{}); outn=b.get('output_name','out.pdf')
    tpl=os.path.join('templates', tid or '')
    if not tpl.endswith('.html') or not os.path.exists(tpl): return jsonify({'ok':False,'error':'html template not found'}),404
    out=os.path.join(tempfile.gettempdir(), outn if outn.endswith('.pdf') else outn+'.pdf')
    with open(tpl,'r',encoding='utf-8') as f: html=Template(f.read()).render(**values)
    HTML(string=html).write_pdf(out); return send_file(out, as_attachment=True, download_name=os.path.basename(out))

if __name__=='__main__':
    app.run(host='0.0.0.0', port=int(os.getenv('PORT','8023')), debug=os.getenv('FLASK_ENV')=='development')
